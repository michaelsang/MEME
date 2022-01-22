from typing import List

import docx as docx

from .QuoteModel import QuoteModel
from .IngestorInterface import IngestorInterface


class DocxIngestor(IngestorInterface):
    """Ingestor class makes use of all other classes.

    Each ingestor has a can_ingest function to validate if it can ingest the file.
    https://knowledge.udacity.com/questions/572306
    https://knowledge.udacity.com/questions/559464
    """
    extensions = ['docx']

    @classmethod
    def parse(cls, path: str) -> List[QuoteModel]:
        if not cls.can_ingest(path):
            raise Exception('cannot ingest exception')

        quotes = []
#        try:
            #docx.Document(path)
        for paragraph in docx.Document(path).paragraphs:
        #line = lines.strip('')
            if paragraph.text != "":
                docx_parse = paragraph.text.split('-')
                docx_quote = QuoteModel(docx_parse[0], docx_parse[1])
                quotes.append(docx_quote)
        return quotes

#        except Exception as e:
#            raise Exception("Docx file parsing error occurred.")

